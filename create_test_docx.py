#!/usr/bin/env python3
"""Create a test DOCX file for VBAM testing"""

from docx import Document

# Create a new document
doc = Document()

# Add title
doc.add_heading('Investment Performance Report (IPR) Documentation', 0)

# Add introduction paragraph
doc.add_heading('Overview', level=1)
p1 = doc.add_paragraph()
p1.add_run('The IPR report analyzes portfolio and benchmark performance and return statistics over selectable time periods. ')
p1.add_run('This comprehensive tool provides detailed insights into investment performance across multiple dimensions.')

# Add inputs section
doc.add_heading('Input Requirements', level=1)
p2 = doc.add_paragraph()
p2.add_run('The IPR system requires the following inputs:')

# Add a list of inputs
inputs = [
    'Portfolio or Composite selection (if Account Type = All)',
    'Marketable Composite (if Account Type = Marketable Composite)', 
    'Return basis selection (Gross or Net Returns)',
    'Time period definition',
    'Reporting frequency (Monthly or Quarterly)'
]

for inp in inputs:
    doc.add_paragraph(inp, style='List Bullet')

# Add outputs section
doc.add_heading('Output Components', level=1)

doc.add_heading('Ret Stats Section', level=2)
p3 = doc.add_paragraph()
p3.add_run('The Ret Stats section provides comprehensive return and risk analysis including:')

ret_stats_items = [
    'Cumulative and Annualized Returns across multiple timeframes',
    'Risk Metrics: Standard Deviation, Volatility measures',
    'Performance Ratios: Sharpe Ratio, Sortino Ratio, Treynor Ratio',
    'Information Ratio and tracking error analysis',
    'Alpha and Beta calculations vs benchmark',
    'R-Squared correlation coefficients',
    'Available timeframes: 3m, 6m, 1-10y, Since Inception'
]

for item in ret_stats_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Ret Stats 2 Section', level=2)
p4 = doc.add_paragraph()
p4.add_run('Extended performance metrics including:')

ret_stats2_items = [
    'Capture Ratios (Up/Down market performance)',
    'Batting Average (percentage of outperformance periods)',
    'Skewness and Kurtosis distribution measures',
    'Risk-adjusted performance indicators'
]

for item in ret_stats2_items:
    doc.add_paragraph(item, style='List Bullet')

doc.add_heading('Additional Output Sections', level=2)
addl_items = [
    'Ret: Monthly or Quarterly performance and excess returns with Risk-Free rate comparison',
    'Trailing Stats: Rolling performance and volatility analysis',
    'Net Asset Values and benchmark comparison metrics'
]

for item in addl_items:
    doc.add_paragraph(item, style='List Bullet')

# Add conclusion
doc.add_heading('Summary', level=1)
p5 = doc.add_paragraph()
p5.add_run('The IPR report serves as a comprehensive performance analysis tool, ')
p5.add_run('providing institutional-grade analytics for portfolio managers and investment professionals. ')
p5.add_run('All metrics can be calculated on gross or net return basis and across flexible time periods.')

# Save the document
doc.save('C:/Projects/enhanced-RAG-3/test_ipr_documentation.docx')
print("Created test_ipr_documentation.docx")
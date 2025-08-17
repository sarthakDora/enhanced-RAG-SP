"""
VBAM Multi-Component RAG Service

Implements support for multiple VBAM components:
- IPR (Investment Performance Report)
- Analytics Report
- Factsheet
- Holdings and Risk

Each component gets its own Qdrant collection with specialized chunking and routing.
"""

from typing import Dict, Any, List, Optional, Tuple
import re
import logging
import uuid
import asyncio
from datetime import datetime
from dataclasses import dataclass
from pathlib import Path
from docx import Document

from qdrant_client.models import Filter, FieldCondition, Match, PointStruct, VectorParams, Distance
from ..models.document import DocumentChunk, DocumentSearchRequest, DocumentSearchResult, DocumentMetadata
from .qdrant_service import QdrantService
from .ollama_service import OllamaService

logger = logging.getLogger(__name__)

# Step 1: Define VBAM Components
VBAM_COMPONENTS = {
    "IPR": "ipr_docs",
    "Analytics Report": "analytics_docs", 
    "Factsheet": "factsheet_docs",
    "Holdings and Risk": "holdings_docs"
}

@dataclass
class VBAMChunk:
    """Enhanced chunk for VBAM components"""
    chunk_id: str
    component: str
    section: str
    content: str
    metadata: Dict[str, Any]
    embedding: Optional[List[float]] = None

class VBAMComponentService:
    """Service for handling VBAM multi-component RAG"""
    
    def __init__(self, qdrant_service: QdrantService, ollama_service: OllamaService):
        self.qdrant = qdrant_service
        self.ollama = ollama_service
        self.embedding_dimension = 768  # nomic-embed-text dimension
        
        # Step 2: Define component collections
        self.component_collections = {
            component: collection_name 
            for component, collection_name in VBAM_COMPONENTS.items()
        }
        
    async def initialize_collections(self) -> bool:
        """Step 2: Create one Qdrant collection per component"""
        try:
            for component, collection_name in self.component_collections.items():
                await self._create_component_collection(collection_name)
                logger.info(f"Initialized collection for {component}: {collection_name}")
            return True
        except Exception as e:
            logger.error(f"Failed to initialize VBAM collections: {e}")
            return False
    
    async def clear_component_collection(self, component: str) -> bool:
        """Clear all documents from a specific component collection"""
        try:
            if component not in self.component_collections:
                logger.error(f"Invalid component: {component}")
                return False
                
            collection_name = self.component_collections[component]
            
            # Delete all points from the collection
            # First get all point IDs
            points = self.qdrant.client.scroll(
                collection_name=collection_name,
                limit=10000  # Get all points
            )[0]
            
            if points:
                point_ids = [point.id for point in points]
                self.qdrant.client.delete(
                    collection_name=collection_name,
                    points_selector=point_ids
                )
                logger.info(f"Deleted {len(point_ids)} points from {collection_name}")
            else:
                logger.info(f"No points to delete from {collection_name}")
            
            logger.info(f"Cleared all documents from {component} collection ({collection_name})")
            return True
            
        except Exception as e:
            logger.error(f"Failed to clear collection for {component}: {e}")
            return False
    
    async def _create_component_collection(self, collection_name: str) -> bool:
        """Create a component-specific collection"""
        try:
            collections = self.qdrant.client.get_collections()
            collection_names = [col.name for col in collections.collections]
            
            if collection_name not in collection_names:
                self.qdrant.client.create_collection(
                    collection_name=collection_name,
                    vectors_config=VectorParams(
                        size=self.embedding_dimension,
                        distance=Distance.COSINE
                    )
                )
                logger.info(f"Created VBAM collection: {collection_name}")
            return True
        except Exception as e:
            logger.error(f"Failed to create collection {collection_name}: {e}")
            return False
    


    def chunk_component_document(self, component: str, content: str, filename: str) -> List[VBAMChunk]:
        """
        Chunk using intelligent content extraction and 70–150 word chunks.
        """
        chunks = []
        
        # Use the improved section extraction method that works with content directly
        try:
            sections = self._extract_vbam_sections(content, component)
            
            # If we got meaningful sections, use them
            if len(sections) > 1:
                cleaned = {k: v for k, v in sections.items() if len(v.split()) >= 10}
            else:
                # fallback to intelligent content splitting
                cleaned = self._intelligent_content_splitting(content, [])

        except Exception as e:
            logger.warning(f"Section extraction failed: {e}, using fallback")
            # fallback to intelligent content splitting
            cleaned = self._intelligent_content_splitting(content, [])

        # now build chunks
        for sec, body in cleaned.items():
            body = body.strip()
            if not body:
                continue
            word_count = len(body.split())
            slug = sec.lower().replace(" ", "_")

            if word_count <= 150:
                cid = f"{component.lower()}_{slug}"
                chunks.append(
                    VBAMChunk(
                        chunk_id=cid,
                        component=component,
                        section=sec,
                        content=body,
                        metadata={
                            "component": component,
                            "section": sec,
                            "chunk_id": cid,
                            "filename": filename,
                            "chunk_index": 0,
                            "word_count": word_count,
                            "created_at": datetime.now().isoformat(),
                        },
                    )
                )
            else:
                parts = self._split_section_intelligently(body, 70, 150)
                for i, p in enumerate(parts):
                    cid = f"{component.lower()}_{slug}_part{i+1}"
                    chunks.append(
                        VBAMChunk(
                            chunk_id=cid,
                            component=component,
                            section=sec,
                            content=p,
                            metadata={
                                "component": component,
                                "section": sec,
                                "chunk_id": cid,
                                "filename": filename,
                                "chunk_index": i,
                                "word_count": len(p.split()),
                                "created_at": datetime.now().isoformat(),
                            },
                        )
                    )

        return chunks

    def chunk_docx_by_headers(self, component: str, file_path: str, filename: str) -> List[VBAMChunk]:
        """
        Chunk DOCX document based on Header 2 styles.
        Each Header 2 becomes a section boundary for chunking.
        """
        chunks = []
        
        try:
            import docx
            doc = docx.Document(file_path)
            
            sections = {}
            current_header = "Overview"  # Default section name
            current_content = []
            
            logger.info(f"Processing docx file: {filename} with {len(doc.paragraphs)} paragraphs")
            
            # First pass: check if this document has reasonable content
            readable_paragraphs = 0
            for para in doc.paragraphs:
                if para.text.strip():
                    test_cleaned = self._clean_docx_text(para.text)
                    if test_cleaned and len(test_cleaned.split()) >= 2:
                        readable_paragraphs += 1
            
            logger.info(f"Found {readable_paragraphs} readable paragraphs out of {len(doc.paragraphs)}")
            
            if readable_paragraphs < 3:
                logger.warning(f"Document {filename} has very few readable paragraphs, may be corrupted")
            
            for para in doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue
                
                # Clean up encoding issues and filter out binary content
                cleaned_text = self._clean_docx_text(text)
                
                # Skip if the text was filtered out as binary content
                if not cleaned_text:
                    logger.debug(f"Filtered out binary content: {text[:100]}...")
                    continue
                elif cleaned_text != text:
                    logger.debug(f"Cleaned text: '{text[:50]}...' -> '{cleaned_text[:50]}...'")
                
                # Additional validation for meaningful content (but allow potential headers)
                if len(cleaned_text.strip()) < 3:
                    # Allow single words that might be headers if they have heading style
                    if 'heading' not in para.style.name.lower():
                        logger.debug(f"Skipping short content (not a heading): '{cleaned_text}'")
                        continue
                    
                # Check if this paragraph is a Header style
                style_name = para.style.name.lower()
                logger.debug(f"Paragraph style: '{style_name}', text: '{cleaned_text[:50]}...'")
                
                # Strict header detection - only actual Heading 2 styles
                is_header = (
                    'heading 2' in style_name or 
                    style_name == 'heading 2' or
                    style_name == 'heading2'
                )
                
                if is_header:
                    # Save previous section if it has content
                    if current_content:
                        section_text = '\n'.join(current_content).strip()
                        if len(section_text.split()) >= 10:  # Minimum word count
                            sections[current_header] = section_text
                            logger.info(f"Found section '{current_header}' with {len(section_text.split())} words")
                        current_content = []
                    
                    # Start new section
                    current_header = cleaned_text
                    logger.info(f"New section header found: {current_header} (style: {style_name})")
                    
                elif 'heading 1' in style_name or style_name == 'heading1':
                    # Also treat Heading 1 as section boundary, but with higher priority
                    if current_content:
                        section_text = '\n'.join(current_content).strip()
                        if len(section_text.split()) >= 10:
                            sections[current_header] = section_text
                            logger.info(f"Found section '{current_header}' with {len(section_text.split())} words")
                        current_content = []
                    
                    current_header = cleaned_text
                    logger.info(f"New H1 section header found: {current_header}")
                    
                else:
                    # Regular content - add to current section (only if it's clean text)
                    if cleaned_text and len(cleaned_text.split()) >= 2:  # Filter out very short fragments
                        current_content.append(cleaned_text)
            
            # Don't forget the last section
            if current_content:
                section_text = '\n'.join(current_content).strip()
                if len(section_text.split()) >= 10:
                    sections[current_header] = section_text
                    logger.info(f"Final section '{current_header}' with {len(section_text.split())} words")
            
            logger.info(f"Extracted {len(sections)} sections from docx: {list(sections.keys())}")
            
            # If no sections found, fallback to intelligent content splitting
            if not sections:
                logger.warning("No header-based sections found, using intelligent content splitting")
                # Collect all cleaned text
                clean_paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        cleaned = self._clean_docx_text(para.text)
                        if cleaned and len(cleaned.split()) >= 2:
                            clean_paragraphs.append(cleaned)
                
                if clean_paragraphs:
                    full_text = '\n'.join(clean_paragraphs)
                    # Use intelligent content splitting instead of single chunk
                    sections = self._intelligent_content_splitting(full_text, ["Overview", "Content", "Details", "Summary"])
            
            # Create chunks from sections
            for section_title, section_content in sections.items():
                section_content = section_content.strip()
                if not section_content:
                    continue
                    
                word_count = len(section_content.split())
                section_slug = section_title.lower().replace(' ', '_').replace('&', 'and')
                
                # If section is within optimal size (70-150 words), create single chunk
                if word_count <= 150:
                    chunk_id = f"{component.lower()}_{section_slug}"
                    chunks.append(VBAMChunk(
                        chunk_id=chunk_id,
                        component=component,
                        section=section_title,
                        content=section_content,
                        metadata={
                            "component": component,
                            "section": section_title,
                            "chunk_id": chunk_id,
                            "filename": filename,
                            "chunk_index": 0,
                            "word_count": word_count,
                            "created_at": datetime.now().isoformat(),
                        }
                    ))
                    logger.info(f"Created single chunk for section '{section_title}' ({word_count} words)")
                    
                else:
                    # Split large sections into multiple chunks
                    sub_chunks = self._split_section_intelligently(section_content, 70, 150)
                    for i, sub_chunk in enumerate(sub_chunks):
                        chunk_id = f"{component.lower()}_{section_slug}_part{i+1}"
                        chunks.append(VBAMChunk(
                            chunk_id=chunk_id,
                            component=component,
                            section=section_title,
                            content=sub_chunk,
                            metadata={
                                "component": component,
                                "section": section_title,
                                "chunk_id": chunk_id,
                                "filename": filename,
                                "chunk_index": i,
                                "word_count": len(sub_chunk.split()),
                                "created_at": datetime.now().isoformat(),
                            }
                        ))
                    logger.info(f"Split section '{section_title}' into {len(sub_chunks)} chunks")
            
            logger.info(f"Created {len(chunks)} total chunks from docx file")
            return chunks
            
        except Exception as e:
            logger.error(f"Failed to process docx by headers: {e}")
            # Fallback to content-based chunking
            try:
                doc = docx.Document(file_path)
                # Clean and filter text properly in fallback mode
                clean_paragraphs = []
                for para in doc.paragraphs:
                    if para.text.strip():
                        cleaned = self._clean_docx_text(para.text)
                        if cleaned and len(cleaned.split()) >= 2:  # Filter out fragments
                            clean_paragraphs.append(cleaned)
                
                full_text = '\n'.join(clean_paragraphs)
                if full_text.strip():
                    return self.chunk_component_document(component, full_text, filename)
                else:
                    logger.warning(f"No clean text extracted from {filename}")
                    return []
            except Exception as e2:
                logger.error(f"Fallback chunking also failed: {e2}")
                return []

    def _clean_docx_text(self, text: str) -> str:
        """Clean up text from DOCX files to fix encoding issues and filter out binary content"""
        if not text:
            return text
        
        # First, check if this looks like binary/encoded content and skip it
        if self._is_binary_content(text):
            logger.debug(f"Skipping binary content: {text[:50]}...")
            return ""
        
        # Replace common problematic Unicode characters
        replacements = {
            '� ': '• ',  # Replace unknown chars with bullet points
            '�': '•',    # Single unknown char to bullet
            '\u2022': '• ',  # Ensure bullet points have space
            '\u2013': '- ',  # En dash to hyphen
            '\u2014': '- ',  # Em dash to hyphen
            '\u201c': '"',   # Left double quote
            '\u201d': '"',   # Right double quote
            '\u2018': "'",   # Left single quote
            '\u2019': "'",   # Right single quote
            '\u00a0': ' ',   # Non-breaking space
        }
        
        # Apply replacements
        cleaned_text = text
        for old, new in replacements.items():
            cleaned_text = cleaned_text.replace(old, new)
        
        # Remove any remaining problematic characters and normalize whitespace
        # Be more aggressive about removing binary-like sequences
        cleaned_text = re.sub(r'[^\w\s\-.,;:!?()\[\]{}"\'/|\\&@#$%^*+=<>~`]', ' ', cleaned_text)
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text).strip()
        
        # Final check - if after cleaning we still have suspicious patterns, return empty
        if (len(cleaned_text) > 20 and 
            (cleaned_text.count(' ') < len(cleaned_text) * 0.1 or  # Very few spaces
             bool(re.search(r'[A-Z]{5,}[a-z]{5,}[A-Z]{5,}', cleaned_text)) or  # Mixed case patterns
             bool(re.search(r'[0-9!@#$%^&*()]{5,}', cleaned_text)))):  # Lots of special chars/numbers
            logger.debug(f"Final filter caught suspicious text: {cleaned_text[:50]}...")
            return ""
        
        return cleaned_text
    
    def _is_binary_content(self, text: str) -> bool:
        """Check if text appears to be binary/encoded content that should be filtered out"""
        if not text or len(text) < 5:
            return False
        
        # First check for the specific pattern you showed: V'Z!:OhBP{(s/BP!Z-T8
        if bool(re.search(r"[VZ]'[VZ]![:\w]*Oh[BP]+\{", text)):
            return True
        
        # Check for patterns that indicate binary content
        binary_indicators = [
            # Common binary patterns
            len([c for c in text if ord(c) < 32 and c not in '\t\n\r']) > len(text) * 0.05,  # Too many control chars (lowered threshold)
            # Base64-like patterns
            bool(re.search(r'^[A-Za-z0-9+/]{15,}={0,2}$', text.strip())),
            # Hex-like patterns
            bool(re.search(r'^[0-9A-Fa-f\s]{15,}$', text.strip())),
            # Random-looking character sequences with mixed special chars
            bool(re.search(r'[!@#$%^&*(){}[\]|\\:";\'<>?,./~`]{8,}', text)),
            # Specific patterns like the one you showed
            bool(re.search(r'[BP]{2,}![VZ\-T\d]+', text)),
            # ZIP/archive signatures
            text.startswith(('PK', '\x50\x4b')),
            # Office document internal structure markers
            bool(re.search(r'(word/|xl/|ppt/|docProps/|_rels/)', text)),
            # XML namespace declarations without proper structure
            bool(re.search(r'^xmlns[^>]{30,}', text.strip())),
            # Very long strings without spaces (likely encoded)
            len(text.split()) == 1 and len(text) > 50,
            # Mixed uppercase/lowercase with many special chars
            bool(re.search(r'[A-Z][a-z]*[!@#$%^&*()]+[A-Z][a-z]*[{}\[\]]+', text)),
            # Check for sequences with many non-ASCII printable characters
            len([c for c in text if ord(c) > 127 or (ord(c) < 32 and c not in '\t\n\r')]) > len(text) * 0.2,
            # Pattern matching multiple curly braces and mixed case
            bool(re.search(r'[{}\[\]]{2,}.*[A-Z]{2,}.*[a-z]{2,}.*[0-9!@#$%^&*()]+', text)),
        ]
        
        return any(binary_indicators)

    def _extract_vbam_sections(self, content: str, component: str) -> Dict[str, str]:
        """
        Updated: improved section extraction.  
        If Heading-2 style markers (e.g. lines that match expected section headers)
        are found, use those as split points. Otherwise use fallback splitting.
        """
        sections = {}
        
        section_headers = {
            "IPR": [
                "Overview", "Navigation", "Inputs", "Outputs", "Ret Stats", "Ret Stats 2",
                "Trailing Stats", "Risk Metrics", "Performance Analysis"
            ],
            "Analytics Report": [
                "Overview", "Navigation", "Inputs", "Outputs", "Factor Attribution",
                "Style Box Analysis", "Portfolio Positioning", "Performance Drivers"
            ],
            "Factsheet": [
                "Overview", "Investment Objective", "Key Statistics",
                "Top Holdings", "Sector Allocation", "Risk Metrics"
            ],
            "Holdings and Risk": [
                "Overview", "Holdings Analysis", "Risk Metrics", "Sector Breakdown",
                "Geographic Allocation", "Risk Assessment", "Concentration Analysis"
            ]
        }
        
        expected_headers = section_headers.get(component, ["Overview", "Content", "Details"])

        lines = content.split("\n")
        current_section = "Overview"
        current_lines = []
        found_headers = 0

        for line in lines:
            stripped = line.strip()
            if not stripped:
                current_lines.append(line)
                continue

            is_header = False
            for header in expected_headers:
                if stripped.lower() == header.lower():
                    # Found header
                    is_header = True
                    if current_lines:
                        sec_content = "\n".join(current_lines).strip()
                        if sec_content:
                            sections[current_section] = sec_content
                    current_section = header
                    current_lines = []
                    found_headers += 1
                    break

            if not is_header:
                current_lines.append(line)

        if current_lines:
            sec_content = "\n".join(current_lines).strip()
            if sec_content:
                sections[current_section] = sec_content

        if found_headers == 0 or len(sections) <= 1:
            # fallback = automatic splitting
            logger.info(f"No explicit headers found; fallback splitting for component {component}")
            return self._intelligent_content_splitting(content, expected_headers)

        return sections


    # def chunk_component_document(self, component: str, content: str, filename: str) -> List[VBAMChunk]:
        """
        Section-based chunking for VBAM documentation following specific rules:
        - 70-150 words per chunk
        - Section-based boundaries
        - Consistent naming scheme
        - Preserve original wording
        """
        chunks = []
        
        # ---- Updated section detection ----
        sections = self._extract_vbam_sections(content, component)

        for section_title, section_content in sections.items():
            section_content = section_content.strip()
            if not section_content or len(section_content.split()) < 20:
                continue

            section_slug = section_title.lower().replace(' ', '_').replace('&', 'and')
            words = section_content.split()

            # ---- Updated splitting method (range-based) ----
            start = 0
            index = 0
            while start < len(words):
                end = min(start + 150, len(words))
                chunk_words = words[start:end]
                # Ensure we have at least 70 words
                if len(chunk_words) < 70 and start != 0:
                    # append the remaining words to the previous chunk
                    prev_chunk = chunks[-1]
                    prev_chunk.content += " " + " ".join(chunk_words)
                    prev_chunk.metadata["word_count"] = len(prev_chunk.content.split())
                    break

                chunk_id = f"{component.lower()}_{section_slug}_part{index+1}" if (start > 0) else f"{component.lower()}_{section_slug}"

                chunk = VBAMChunk(
                    chunk_id=chunk_id,
                    component=component,
                    section=section_title,
                    content=" ".join(chunk_words),
                    metadata={
                        "component": component,
                        "section": section_title,
                        "chunk_id": chunk_id,
                        "filename": filename,
                        "chunk_index": index,
                        "word_count": len(chunk_words),
                        "created_at": datetime.now().isoformat()
                    }
                )
                chunks.append(chunk)
                start += 150
                index += 1

        logger.info(f"Created {len(chunks)} chunks for {component} documentation")
        return chunks

    def _intelligent_content_splitting(self, content: str, expected_headers: List[str]) -> Dict[str, str]:
        """
        Intelligent content splitting when no clear headers are found
        """
        sections = {}
        
        # Split by paragraphs and group intelligently
        paragraphs = [p.strip() for p in content.split('\n\n') if p.strip()]
        
        # Also try splitting by single newlines if no double-newline paragraphs
        if len(paragraphs) <= 1:
            paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        
        logger.info(f"Intelligent splitting: Found {len(paragraphs)} paragraphs from content")
        
        if len(paragraphs) <= 3:
            # Short document - split by word count instead of paragraphs
            words = content.split()
            if len(words) > 300:  # Large document, split into 4 parts
                quarter = len(words) // 4
                sections["Overview"] = ' '.join(words[:quarter])
                sections["Navigation"] = ' '.join(words[quarter:2*quarter])
                sections["Inputs"] = ' '.join(words[2*quarter:3*quarter])
                sections["Outputs"] = ' '.join(words[3*quarter:])
            elif len(words) > 150:  # Medium document, split into 3 parts
                third = len(words) // 3
                sections["Overview"] = ' '.join(words[:third])
                sections["Details"] = ' '.join(words[third:2*third])
                sections["Summary"] = ' '.join(words[2*third:])
            elif len(words) > 70:  # Small document, split in half
                half = len(words) // 2
                sections["Overview"] = ' '.join(words[:half])
                sections["Details"] = ' '.join(words[half:])
            else:
                sections["Overview"] = content
        else:
            # Distribute paragraphs across expected sections
            if len(paragraphs) >= len(expected_headers):
                section_size = max(1, len(paragraphs) // len(expected_headers))
                
                for i, header in enumerate(expected_headers):
                    start_idx = i * section_size
                    end_idx = start_idx + section_size if i < len(expected_headers) - 1 else len(paragraphs)
                    
                    if start_idx < len(paragraphs):
                        section_content = '\n\n'.join(paragraphs[start_idx:end_idx])
                        if section_content and len(section_content.split()) >= 5:
                            sections[header] = section_content
                            logger.info(f"Created intelligent section '{header}' with {len(section_content)} chars")
            else:
                # Fewer paragraphs than expected headers, distribute differently
                for i, paragraph in enumerate(paragraphs):
                    if i < len(expected_headers) and len(paragraph.split()) >= 5:
                        sections[expected_headers[i]] = paragraph
                        logger.info(f"Assigned paragraph to '{expected_headers[i]}' with {len(paragraph)} chars")
        
        logger.info(f"Intelligent splitting created {len(sections)} sections: {list(sections.keys())}")
        return sections
    
    def _split_section_intelligently(self, content: str, min_words: int, max_words: int) -> List[str]:
        """
        Split a long section into chunks while preserving logical boundaries
        """
        chunks = []
        
        # Split by sentences and paragraphs
        sentences = re.split(r'(?<=[.!?])\s+', content)
        current_chunk = []
        current_word_count = 0
        
        for sentence in sentences:
            sentence_words = len(sentence.split())
            
            # If adding this sentence would exceed max_words, save current chunk
            if current_word_count + sentence_words > max_words and current_word_count >= min_words:
                if current_chunk:
                    chunks.append(' '.join(current_chunk))
                    current_chunk = [sentence]
                    current_word_count = sentence_words
                else:
                    # Single sentence is too long, add it anyway
                    chunks.append(sentence)
                    current_word_count = 0
            else:
                current_chunk.append(sentence)
                current_word_count += sentence_words
        
        # Add remaining content
        if current_chunk:
            chunks.append(' '.join(current_chunk))
        
        return chunks
    
    def _chunk_section_content(self, content: str, min_words: int, max_words: int) -> List[str]:
        """Chunk section content into optimal sizes"""
        words = content.split()
        chunks = []
        current_chunk = []
        
        for word in words:
            current_chunk.append(word)
            
            # Check if we should create a chunk
            if len(current_chunk) >= max_words:
                # Look for good break point
                chunk_text = ' '.join(current_chunk)
                break_point = self._find_break_point(chunk_text, min_words)
                
                if break_point > 0:
                    chunk = ' '.join(current_chunk[:break_point])
                    chunks.append(chunk)
                    current_chunk = current_chunk[break_point:]
                else:
                    chunks.append(chunk_text)
                    current_chunk = []
        
        # Add remaining content
        if current_chunk and len(current_chunk) >= min_words // 2:
            chunks.append(' '.join(current_chunk))
        
        return chunks if chunks else [content]
    
    def _find_break_point(self, text: str, min_words: int) -> int:
        """Find good break point in text (sentence boundary)"""
        words = text.split()
        if len(words) <= min_words:
            return 0
        
        # Look for sentence endings after minimum words
        for i in range(min_words, len(words)):
            word = words[i]
            if word.endswith(('.', '!', '?')) and i < len(words) - 1:
                return i + 1
        
        # Fallback to word boundary
        return len(words) // 2
    
    async def store_component_chunks(self, chunks: List[VBAMChunk]) -> bool:
        """Step 4 & 5: Embed and store chunks in appropriate collections"""
        try:
            if not chunks:
                return True
            
            # Generate embeddings for all chunks
            texts = [chunk.content for chunk in chunks]
            embeddings = await self.ollama.generate_embeddings(texts)
            
            if not embeddings or len(embeddings) != len(texts):
                raise RuntimeError(f"Embedding mismatch: {len(embeddings)} vs {len(texts)}")
            
            # Add embeddings to chunks
            for chunk, embedding in zip(chunks, embeddings):
                chunk.embedding = [float(x) for x in embedding]
            
            # Group chunks by component
            component_chunks = {}
            for chunk in chunks:
                component = chunk.component
                if component not in component_chunks:
                    component_chunks[component] = []
                component_chunks[component].append(chunk)
            
            # Store each component's chunks in its collection
            for component, comp_chunks in component_chunks.items():
                collection_name = self.component_collections.get(component)
                if not collection_name:
                    logger.warning(f"Unknown component: {component}")
                    continue
                
                await self._store_chunks_in_collection(comp_chunks, collection_name)
            
            return True
            
        except Exception as e:
            logger.error(f"Failed to store component chunks: {e}")
            return False
    
    async def _store_chunks_in_collection(self, chunks: List[VBAMChunk], collection_name: str):
        """Store chunks in specific collection"""
        points = []
        
        for chunk in chunks:
            if not chunk.embedding:
                logger.warning(f"Chunk {chunk.chunk_id} has no embedding, skipping")
                continue
            
            payload = {
                "content": chunk.content,
                "component": chunk.component,
                "section": chunk.section,
                "chunk_id": chunk.chunk_id,
                **chunk.metadata
            }
            
            point = PointStruct(
                id=str(uuid.uuid4()),
                vector=chunk.embedding,
                payload=payload
            )
            points.append(point)
        
        if points:
            self.qdrant.client.upsert(
                collection_name=collection_name,
                points=points
            )
            logger.info(f"Stored {len(points)} chunks in {collection_name}")
    
    def route_question_to_component(self, question: str) -> Optional[str]:
        """Step 6: Simple keyword-based router"""
        question_lower = question.lower()
        
        # Component keywords mapping
        component_keywords = {
            "IPR": ["ipr", "investment performance", "ret stats", "performance report", 
                   "return", "sharpe", "sortino", "alpha", "beta", "trailing"],
            "Analytics Report": ["analytics", "factor attribution", "style box", 
                               "positioning", "factor exposure", "factor contribution"],
            "Factsheet": ["factsheet", "fact sheet", "investment objective", "top holdings",
                         "sector allocation", "key statistics"],
            "Holdings and Risk": ["holdings", "risk", "portfolio composition", 
                                "concentration", "credit quality", "geographic"]
        }
        
        # Score each component
        component_scores = {}
        for component, keywords in component_keywords.items():
            score = 0
            for keyword in keywords:
                if keyword in question_lower:
                    score += 1
            component_scores[component] = score
        
        # Return component with highest score
        if component_scores:
            best_component = max(component_scores, key=component_scores.get)
            if component_scores[best_component] > 0:
                return best_component
        
        return None
    
    async def search_component(self, component: str, question: str, top_k: int = 10) -> List[Dict[str, Any]]:
        """Search within a specific component collection"""
        try:
            collection_name = self.component_collections.get(component)
            if not collection_name:
                logger.warning(f"Unknown component: {component}")
                return []
            
            # Check if collection exists
            collections = self.qdrant.client.get_collections()
            collection_names = [col.name for col in collections.collections]
            if collection_name not in collection_names:
                logger.warning(f"Collection {collection_name} does not exist")
                return []
            
            # Generate query embedding
            query_embedding = await self.ollama.generate_embedding(question)
            
            # Search collection
            search_results = self.qdrant.client.search(
                collection_name=collection_name,
                query_vector=query_embedding,
                limit=top_k,
                with_payload=True
            )
            
            # Format results
            results = []
            for result in search_results:
                results.append({
                    "content": result.payload.get("content", ""),
                    "component": result.payload.get("component", ""),
                    "section": result.payload.get("section", ""),
                    "score": result.score,
                    "metadata": result.payload
                })
            
            return results
            
        except Exception as e:
            logger.error(f"Failed to search component {component}: {e}")
            return []
    
    async def answer_component_question(self, question: str, component: Optional[str] = None, conversation_history: Optional[List] = None) -> Dict[str, Any]:
        """Step 7: Main retrieval + prompt system"""
        try:
            # Check if this is a conversation summary request first (before component routing)
            is_summary_request = self._is_summary_request(question)
            logger.info(f"VBAM: Processing question='{question}', is_summary={is_summary_request}, has_history={bool(conversation_history)}")
            
            if is_summary_request:
                if conversation_history and len(conversation_history) > 0:
                    logger.info(f"VBAM: Building summary prompt with {len(conversation_history)} messages")
                    prompt = self._build_summary_prompt(conversation_history, question)
                    system_prompt = """You are a VBAM support specialist creating conversation summaries. 
Provide clear, structured summaries of VBAM-related discussions, highlighting key questions asked and answers provided."""
                    
                    # Generate summary response
                    ollama_response = await self.ollama.generate_response(
                        prompt=prompt,
                        context="",  # Context is included in the prompt
                        temperature=0.1,
                        system_prompt=system_prompt
                    )
                    
                    # Extract just the response text from the ollama response dictionary
                    response_text = ollama_response.get("response", "Summary generation failed") if isinstance(ollama_response, dict) else str(ollama_response)
                    
                    return {
                        "response": response_text,
                        "component": "Summary",
                        "results_count": len(conversation_history),
                        "sections_referenced": ["Conversation History"],
                        "routed": True
                    }
                else:
                    # No conversation history to summarize
                    return {
                        "response": "I don't have any previous conversation to summarize. Please ask me questions about VBAM components (IPR, Analytics Report, Factsheet, Holdings and Risk) and I'll be happy to help and maintain context for future summary requests.",
                        "component": None,
                        "results_count": 0,
                        "routed": False
                    }
            
            # Route to component if not specified (non-summary requests)
            if not component:
                component = self.route_question_to_component(question)
                
                # If no component detected from current question, try to infer from conversation history
                if not component and conversation_history:
                    component = self._infer_component_from_history(conversation_history, question)
                    if component:
                        logger.info(f"VBAM: Inferred component '{component}' from conversation history")
                
                if not component:
                    return {
                        "response": "Can you clarify which component your question refers to? Available components: IPR, Analytics Report, Factsheet, Holdings and Risk.",
                        "component": None,
                        "routed": False
                    }
            
            # Search component
            search_results = await self.search_component(component, question)
            
            if not search_results:
                return {
                    "response": f"No relevant information found in {component} documentation.",
                    "component": component,
                    "results_count": 0
                }
            
            # Build context from search results with better formatting
            context_parts = []
            for i, result in enumerate(search_results, 1):
                section = result["section"]
                content = result["content"]
                score = result.get("score", 0.0)
                context_parts.append(f"SECTION {i}: {section} (Relevance: {score:.2f})\n{content}\n")
            
            context = "\n".join(context_parts)
            
            # Generate response with component-specific prompt (summary requests already handled above)
            prompt = self._build_component_prompt(component, question, context, conversation_history)
            # Enhanced system prompt for VBAM
            system_prompt = f"""You are a specialized VBAM (Virtus Business Application Manager) support expert for the {component} component. 
            
You help financial professionals understand VBAM features, navigation, inputs, outputs, and functionality. 
Provide clear, actionable guidance based on the official VBAM documentation.

When provided with conversation history, use it to understand context and maintain continuity in the discussion."""
            
            ollama_response = await self.ollama.generate_response(
                prompt=prompt,
                context="",  # Context is included in the prompt
                system_prompt=system_prompt,
                temperature=0.1
            )
            
            # Extract just the response text from the ollama response dictionary
            response_text = ollama_response.get("response", "Response generation failed") if isinstance(ollama_response, dict) else str(ollama_response)
            
            return {
                "response": response_text,
                "component": component,
                "results_count": len(search_results),
                "sections_referenced": list(set(r["section"] for r in search_results)),
                "routed": component != component  # True if auto-routed
            }
            
        except Exception as e:
            logger.error(f"Failed to answer component question: {e}")
            return {
                "response": f"Error processing question: {str(e)}",
                "component": component,
                "error": True
            }
    
    def _build_component_prompt(self, component: str, question: str, context: str, conversation_history: Optional[List] = None) -> str:
        """Build component-specific prompt with optional conversation history"""
        prompt_parts = [f"You are an expert VBAM product support specialist helping users understand {component} functionality."]
        
        # Add conversation history if provided
        if conversation_history and len(conversation_history) > 0:
            prompt_parts.append("\nCONVERSATION HISTORY:")
            for msg in conversation_history:
                role = msg.get('role', 'unknown')
                content = msg.get('content', '')
                if role == 'user':
                    prompt_parts.append(f"User: {content}")
                elif role == 'assistant':
                    # Truncate long assistant responses for context
                    truncated_content = content[:200] + "..." if len(content) > 200 else content
                    prompt_parts.append(f"Assistant: {truncated_content}")
            prompt_parts.append("")
        
        prompt_parts.extend([
            f"VBAM COMPONENT: {component}",
            "",
            "DOCUMENTATION CONTEXT:",
            context,
            "",
            f"USER QUESTION: {question}",
            "",
            "INSTRUCTIONS:",
            "1. Answer the user's question based on the provided documentation context",
            "2. Use the conversation history to understand context and avoid repeating information unnecessarily",
            "3. If the user refers to something mentioned earlier (like 'it', 'that component', etc.), use the conversation history to understand what they mean",
            "4. Be specific and helpful - include exact details from the documentation",
            "5. If the question asks about specific features, inputs, outputs, or navigation, provide step-by-step information",
            "6. Reference the specific sections where you found the information",
            "7. If the context doesn't contain enough information to fully answer the question, say so clearly",
            "8. Use a professional but friendly tone appropriate for financial software support",
            "9. Include practical usage tips when relevant",
            "",
            "ANSWER:"
        ])
        
        return "\n".join(prompt_parts)

    def _build_summary_prompt(self, conversation_history: List, question: str) -> str:
        """Build prompt for conversation summary requests"""
        prompt_parts = [
            "Please provide a comprehensive summary of our VBAM-related conversation.",
            "",
            "CONVERSATION HISTORY:"
        ]
        
        # Add conversation history with better formatting
        conversation_count = 0
        for i, msg in enumerate(conversation_history):
            role = msg.get('role', 'unknown')
            content = msg.get('content', '')
            
            if role == 'user':
                conversation_count += 1
                prompt_parts.append(f"\n=== Conversation {conversation_count} ===")
                prompt_parts.append(f"User Question: {content}")
            elif role == 'assistant':
                # Limit assistant response length for summary context
                truncated_content = content[:300] + "..." if len(content) > 300 else content
                prompt_parts.append(f"Assistant Response: {truncated_content}")
        
        prompt_parts.extend([
            "",
            f"USER SUMMARY REQUEST: {question}",
            "",
            "INSTRUCTIONS:",
            "1. Create a well-structured summary with clear sections",
            "2. Use the following format:",
            "   ## Conversation Summary",
            "   ### VBAM Components Discussed:",
            "   ### Key Topics Covered:",
            "   ### Main Questions & Answers:",
            "   ### Important Details:",
            "3. Group related questions and topics together",
            "4. Highlight specific VBAM components (IPR, Analytics Report, Factsheet, Holdings and Risk) discussed",
            "5. Include the main questions asked and key answers provided",
            "6. Use bullet points and clear formatting for readability",
            "7. Focus on actionable information and important VBAM functionality details",
            "8. If no meaningful conversation exists, say so clearly",
            "",
            "SUMMARY:"
        ])
        
        return "\n".join(prompt_parts)

    def _is_summary_request(self, question: str) -> bool:
        """Check if the user is asking for a conversation summary"""
        summary_keywords = [
            'summarize', 'summary', 'recap', 'review', 'what did we discuss',
            'conversation summary', 'what have we talked about', 'overview of our chat',
            'what questions did i ask', 'what topics did we cover', 'give me a summary',
            'can you summarize', 'please summarize', 'summarize our conversation',
            'sum up', 'wrap up', 'overview', 'what have we covered'
        ]
        question_lower = question.lower().strip()
        
        # Check for exact matches and partial matches
        is_summary = any(keyword in question_lower for keyword in summary_keywords)
        
        # Additional pattern matching for summary requests
        summary_patterns = [
            'summarize our',
            'sum up our',
            'what did we',
            'what have we',
            'overview of',
            'recap of'
        ]
        is_summary = is_summary or any(pattern in question_lower for pattern in summary_patterns)
        
        logger.info(f"VBAM Summary Detection: Question='{question}' -> Is Summary: {is_summary}")
        return is_summary

    def _infer_component_from_history(self, conversation_history: List, current_question: str) -> Optional[str]:
        """Infer VBAM component from conversation history for context awareness"""
        if not conversation_history:
            return None
        
        # Look for component mentions in recent conversation (last 6 messages)
        recent_messages = conversation_history[-6:] if len(conversation_history) > 6 else conversation_history
        
        component_mentions = {}
        
        for msg in reversed(recent_messages):  # Start from most recent
            content = msg.get('content', '').lower()
            
            # Look for explicit component mentions
            if 'ipr' in content or 'investment performance report' in content:
                component_mentions['IPR'] = component_mentions.get('IPR', 0) + 3
            if 'analytics' in content and 'report' in content:
                component_mentions['Analytics Report'] = component_mentions.get('Analytics Report', 0) + 3
            if 'factsheet' in content or 'fact sheet' in content:
                component_mentions['Factsheet'] = component_mentions.get('Factsheet', 0) + 3
            if 'holdings' in content and 'risk' in content:
                component_mentions['Holdings and Risk'] = component_mentions.get('Holdings and Risk', 0) + 3
            
            # Look for component-specific keywords
            if any(word in content for word in ['sharpe', 'sortino', 'alpha', 'beta', 'return', 'ret stats']):
                component_mentions['IPR'] = component_mentions.get('IPR', 0) + 1
            if any(word in content for word in ['factor attribution', 'style box', 'factor exposure']):
                component_mentions['Analytics Report'] = component_mentions.get('Analytics Report', 0) + 1
            if any(word in content for word in ['top holdings', 'sector allocation', 'key statistics']):
                component_mentions['Factsheet'] = component_mentions.get('Factsheet', 0) + 1
            if any(word in content for word in ['portfolio composition', 'concentration', 'credit quality']):
                component_mentions['Holdings and Risk'] = component_mentions.get('Holdings and Risk', 0) + 1
        
        # Check if current question contains contextual references
        current_lower = current_question.lower()
        contextual_refs = ['it', 'that', 'this', 'outputs', 'inputs', 'features', 'same', 'also']
        has_contextual_ref = any(ref in current_lower for ref in contextual_refs)
        
        # If we have contextual references and component mentions, return the most mentioned component
        if has_contextual_ref and component_mentions:
            best_component = max(component_mentions, key=component_mentions.get)
            logger.info(f"VBAM Context: Inferred '{best_component}' from history. Scores: {component_mentions}")
            return best_component
        
        return None
    
    async def get_component_stats(self) -> Dict[str, Any]:
        """Get statistics for all component collections"""
        stats = {}
        
        for component, collection_name in self.component_collections.items():
            try:
                collections = self.qdrant.client.get_collections()
                collection_names = [col.name for col in collections.collections]
                
                if collection_name in collection_names:
                    info = self.qdrant.client.get_collection(collection_name)
                    stats[component] = {
                        "collection_name": collection_name,
                        "points_count": info.points_count,
                        "indexed_vectors_count": info.indexed_vectors_count,
                        "status": info.status
                    }
                else:
                    stats[component] = {
                        "collection_name": collection_name,
                        "points_count": 0,
                        "indexed_vectors_count": 0,
                        "status": "not_exists"
                    }
            except Exception as e:
                logger.error(f"Failed to get stats for {component}: {e}")
                stats[component] = {
                    "collection_name": collection_name,
                    "error": str(e)
                }
        
        return stats
    
    async def process_component_document(self, component: str, content: str = None, filename: str = None, file_path: str = None) -> bool:
        """
        Process a document for a specific component.
        Can handle either content string or file_path for docx files.
        """
        try:
            if component not in self.component_collections:
                raise ValueError(f"Unknown component: {component}")

            # ✅ Chunk using docx-aware chunker if we have a file path, otherwise use content
            if file_path and filename and filename.lower().endswith('.docx'):
                chunks = self.chunk_docx_by_headers(component, file_path, filename)
            elif content:
                chunks = self.chunk_component_document(component, content, filename)
            else:
                raise ValueError("Either content or file_path must be provided")

            if not chunks:
                logger.warning(f"No chunks generated for {filename}")
                return False

            success = await self.store_component_chunks(chunks)
            logger.info(
                f"Processed {filename} for {component}: {len(chunks)} chunks, success={success}"
            )
            return success

        except Exception as e:
            logger.error(f"Failed to process document for {component}: {e}")
            return False